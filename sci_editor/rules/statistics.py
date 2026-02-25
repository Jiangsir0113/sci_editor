"""规则 3.22 — 统计符号斜体检查与修复

正文中统计符号须为斜体：P, Z, F, t, r, n, df, χ, U 等
"""

import re
from docx.oxml.ns import qn
from ..models import Issue, Severity, SectionType, DocumentStructure
from ..engine import BaseRule


# 需要斜体的统计单字母符号
STAT_SYMBOLS = {
    "P", "Z", "F", "t", "r", "n", "df", "χ", "U",
}

# 匹配模式：符号后跟 空格+运算符 或 = 或 < 或 > 即为统计符号
STAT_PATTERN = re.compile(
    r"(?<!\w)(P|Z|F|t|r|n|df|χ|U)(?=\s*[=<>≤≥])",
)

# 额外模式：χ² test, U test 等
STAT_TEST_PATTERN = re.compile(
    r"(?<!\w)(χ|U)(?=\s*(?:²\s*)?(?:test|检验))",
    re.IGNORECASE,
)


def _run_is_italic(run) -> bool:
    return bool(run.italic)


def _split_and_italicize(paragraph, sym, start_pos_in_para):
    """
    在段落中找到指定位置的符号，将其拆分为独立的斜体 Run。
    """
    pos = 0
    sym_len = len(sym)
    for i, run in enumerate(paragraph.runs):
        run_text = run.text
        run_len = len(run_text)
        
        if pos <= start_pos_in_para < pos + run_len:
            start_in_run = start_pos_in_para - pos
            end_in_run = start_in_run + sym_len
            
            # 基础文本
            before_text = run_text[:start_in_run]
            sym_text = run_text[start_in_run:end_in_run]
            after_text = run_text[end_in_run:]
            
            # 如果符号已经是单独的 Run 且已经是斜体，跳过
            if not before_text and not after_text and run.italic:
                return False

            # 1. 更新当前 Run 为 before_text
            run.text = before_text
            # 2. 插入符号 Run
            sym_run = paragraph.add_run(sym_text)
            # 3. 插入剩余文本 Run
            after_run = paragraph.add_run(after_text)

            # 格式保留：复制原始 Run 的其他属性
            from ..utils import copy_run_format
            copy_run_format(run, sym_run)
            copy_run_format(run, after_run)

            # 应用特殊格式
            sym_run.italic = True
            # after_run.italic 已在 copy_run_format 中从 run 继承
            
            # 手动调整 XML 顺序
            p = paragraph._p
            r_orig = run._r
            r_sym = sym_run._r
            r_after = after_run._r
            
            p.insert(p.index(r_orig) + 1, r_sym)
            p.insert(p.index(r_sym) + 1, r_after)
            
            return True
                
        pos += run_len
    return False


class StatisticsItalicRule(BaseRule):
    rule_id = "3.22"
    rule_name = "统计符号斜体"
    section_name = "3.22 统计符号"

    def check(self, doc: DocumentStructure):
        issues = []
        body = doc.get_section(SectionType.BODY)
        abstract = doc.get_section(SectionType.ABSTRACT)
        tables = doc.get_section(SectionType.TABLES)

        for section in [s for s in [body, abstract, tables] if s]:
            for idx, para in section.paragraphs:
                text = para.text
                for pattern in [STAT_PATTERN, STAT_TEST_PATTERN]:
                    for m in pattern.finditer(text):
                        sym = m.group(1)
                        # 检查对应位置的 Run 是否为斜体
                        pos = 0
                        is_italic = False
                        for run in para.runs:
                            if pos <= m.start() < pos + len(run.text):
                                is_italic = _run_is_italic(run)
                                break
                            pos += len(run.text)
                        
                        if not is_italic:
                            issues.append(Issue(
                                rule_id=self.rule_id,
                                rule_name=self.rule_name,
                                severity=Severity.WARNING,
                                message=f"统计符号 '{sym}' 应为斜体",
                                section=section.section_type.value,
                                paragraph_index=idx,
                                context=text[max(0, m.start()-10):m.end()+10],
                                fixable=True,
                                suggestion=f"将 '{sym}' 设为斜体",
                            ))
        return issues

    def fix(self, doc: DocumentStructure, issue: Issue) -> bool:
        idx = issue.paragraph_index
        if idx < 0 or idx >= len(doc.all_paragraphs):
            return False

        para = doc.all_paragraphs[idx]
        text = para.text
        changed = False
        
        for pattern in [STAT_PATTERN, STAT_TEST_PATTERN]:
            for m in pattern.finditer(text):
                if _split_and_italicize(para, m.group(1), m.start()):
                    changed = True
        
        if changed:
            issue.fix_description = "统计符号已修正为斜体"
        return changed


