"""规则 — 页脚处理 (页码 / 总页数)"""

import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from ..models import Issue, Severity, SectionType, DocumentStructure
from ..engine import BaseRule

class PageFooterRule(BaseRule):
    rule_id = "3.31"
    rule_name = "页脚页码格式"
    section_name = "3.31 页脚"

    def _add_field(self, run, field_type):
        """在 run 中插入 Word 域 (PAGE 或 NUMPAGES)"""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = field_type

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'separate')

        t = OxmlElement('w:t')
        t.text = "0"  # 占位符

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(ns.qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(t)
        run._r.append(fldChar3)

    def check(self, doc: DocumentStructure):
        issues = []
        # 检查是否所有 section 都有符合要求的页脚
        has_correct_footer = True
        for section in doc.word_doc.sections:
            footer_text = "".join(p.text for p in section.footer.paragraphs).strip()
            # 由于域代码在静态检查时可能无法正确读取数值，我们主要检查是否存在 "/" 结构
            if "/" not in footer_text:
                has_correct_footer = False
                break
        
        if not has_correct_footer:
            issues.append(Issue(
                rule_id=self.rule_id,
                rule_name=self.rule_name,
                severity=Severity.INFO,
                message="文档缺少 '页码 / 总页数' 格式的页脚",
                section=self.section_name,
                fixable=True
            ))
        return issues

    def fix(self, doc: DocumentStructure, issue: Issue) -> bool:
        # 1. 获取目标字体 (采样正文第一段)
        target_font_name = "Times New Roman"
        target_font_size = None
        
        body = doc.get_section(SectionType.BODY)
        if body and body.paragraphs:
            first_para = body.paragraphs[0][1]
            if first_para.runs:
                r = first_para.runs[0]
                if r.font.name: target_font_name = r.font.name
                if r.font.size: target_font_size = r.font.size

        # 2. 遍历所有 section 添加页脚
        for section in doc.word_doc.sections:
            # 清除旧页脚内容
            for p in section.footer.paragraphs:
                p.text = ""
            
            # 如果没有段落，新建一个
            if not section.footer.paragraphs:
                footer_para = section.footer.add_paragraph()
            else:
                footer_para = section.footer.paragraphs[0]
            
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 插入：[PAGE] / [NUMPAGES]
            run_page = footer_para.add_run()
            self._add_field(run_page, 'PAGE')
            
            run_sep = footer_para.add_run(" / ")
            
            run_total = footer_para.add_run()
            self._add_field(run_total, 'NUMPAGES')
            
            # 应用字体
            for run in footer_para.runs:
                run.font.name = target_font_name
                if target_font_size:
                    run.font.size = target_font_size
        
        issue.fix_description = f"已添加居中页脚 '页码 / 总页数'，使用字体 {target_font_name}"
        return True
