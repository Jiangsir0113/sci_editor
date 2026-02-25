"""文档解析器：将 .docx 拆解为结构化区段"""

import re
from docx import Document
from .models import SectionType, Section, DocumentStructure


# 区段识别关键词映射
SECTION_KEYWORDS = {
    SectionType.CONTRIBUTIONS: [
        r"^author\s+contributions?\s*:",
        r"^作者贡献\s*[:：]",
    ],
    SectionType.FUNDING: [
        r"^(fund(ing|ed\s+by)|support(ed)?\s+by)\s*:",
        r"^基金\s*[:：]",
        r"^supported\s+by\b",
    ],
    SectionType.CORRESPONDING: [
        r"^corresponding\s+author\s*:",
        r"^通讯作者\s*[:：]",
    ],
    SectionType.ABSTRACT: [
        r"^abstract\s*[:：]?$",
        r"^摘要\s*[:：]?$",
    ],
    SectionType.CORE_TIP: [
        r"^core\s+tip\s*[:：]?",
    ],
    SectionType.KEYWORDS: [
        r"^key\s*words?\s*[:：]",
        r"^关键词\s*[:：]",
    ],
    SectionType.REFERENCES: [
        r"^references?\s*$",
        r"^参考文献\s*$",
    ],
    SectionType.FOOTNOTES: [
        r"^footnote[s]?\s*[:：]?$",
        r"^脚注\s*[:：]?$",
        r"^(article|manuscript)\s+highlights?\s*[:：]?$",
    ],
}

# 正文起始标志
BODY_START_KEYWORDS = [
    r"^introduction\s*$",
    r"^methods\s*$",
    r"^(materials\s+and\s+)?methods\s*$",
    r"^statistical\s+analysis\s*$",
    r"^引言\s*$",
    r"^前言\s*$",
    r"^1[\.\s]",
]

# 正文结束标志（参考文献开始）
BODY_END_KEYWORDS = [
    r"^references?\s*$",
    r"^参考文献\s*$",
]

# WJG/Baishideng 文档头元信息关键词（用于跳过）
WJG_META_PATTERNS = [
    r"^name\s+of\s+journal\s*:",
    r"^manuscript\s+no\s*:",
    r"^manuscript\s+type\s*:",
    r"^received\s*:",
    r"^revised\s*:",
    r"^accepted\s*:",
    r"^published\s+online\s*:",
    r"^peer-review\s+started\s*:",
]


def _classify_paragraph(text: str) -> SectionType:
    """根据段落文本识别区段类型"""
    text_stripped = text.strip()
    if not text_stripped:
        return SectionType.UNKNOWN

    # 识别 Figure/Table 标注
    if re.match(r"^(Figure|Fig\.?)\s+\d+", text_stripped, re.IGNORECASE):
        return SectionType.FIGURES
    if re.match(r"^Table\s+\d+", text_stripped, re.IGNORECASE):
        return SectionType.TABLES

    # 遍历关键词映射
    for section_type, patterns in SECTION_KEYWORDS.items():
        for pattern in patterns:
            if re.search(pattern, text_stripped, re.IGNORECASE):
                return section_type

    return SectionType.UNKNOWN


def _is_wjg_meta(text: str) -> bool:
    """判断是否为 WJG/Baishideng 文档元信息行"""
    for pattern in WJG_META_PATTERNS:
        if re.match(pattern, text.strip(), re.IGNORECASE):
            return True
    return False


def _is_affiliation_line(text: str) -> bool:
    """判断是否为机构行"""
    return bool(re.search(
        r"(Department\s+of|University|Hospital|Institute|College|School\s+of|Center|Centre|Laboratory|Faculty)\b",
        text, re.IGNORECASE
    ))


def _is_author_line(text: str) -> bool:
    """判断是否为作者列表行"""
    text = text.strip()
    if not text:
        return False
    # 作者列表特征：多个人名用逗号分隔，大部分是大写字母开头的短词
    parts = [p.strip() for p in text.split(",")]
    if len(parts) < 2:
        return False
    # 检查是否大部分看起来像人名（2-3个词，首字母大写）
    name_count = 0
    for part in parts:
        words = part.strip().split()
        if 1 <= len(words) <= 5 and all(
            w[0].isupper() or w in ("de", "van", "von", "da", "di", "el", "al", "bin")
            for w in words if len(w) > 1
        ):
            name_count += 1
    return name_count >= 2 and name_count / len(parts) > 0.5


def parse_document(filepath: str) -> DocumentStructure:
    """
    解析 .docx 文件为结构化文档。
    """
    doc = Document(filepath)
    ds = DocumentStructure(filepath=filepath)
    ds.word_doc = doc
    ds.all_paragraphs = list(doc.paragraphs)
    ds.tables = list(doc.tables)

    # 提取 Manuscript Type
    for para in doc.paragraphs[:10]:
        m = re.match(r"manuscript\s+type\s*:\s*(.+)", para.text.strip(), re.IGNORECASE)
        if m:
            ds.manuscript_type = m.group(1).strip().upper()
            break

    paragraphs = list(doc.paragraphs)

    # 检测是否为 WJG 格式
    is_wjg = False
    for i in range(min(5, len(paragraphs))):
        if _is_wjg_meta(paragraphs[i].text):
            is_wjg = True
            break

    para_labels = []
    current_section = SectionType.UNKNOWN
    in_body = False
    in_references = False
    in_abstract = False
    found_title = False
    found_running_title = False
    found_authors = False
    found_affiliations = False

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            para_labels.append((i, current_section, para))
            continue

        detected = _classify_paragraph(text)

        # 状态切换逻辑
        if detected == SectionType.ABSTRACT:
            in_abstract = True
            current_section = SectionType.ABSTRACT
        elif detected == SectionType.KEYWORDS:
            in_abstract = False
            current_section = SectionType.KEYWORDS
        elif in_abstract:
            # 在摘要开始后，直到遇到 Keywords 标志或正式的正文标志前，都属于 ABSTRACT
            # 这里我们通过 Keywords 来显式关闭摘要区段。
            # 如果解析到了明确的区段（如 CORE TITP），虽然 current_section 会变，但 in_abstract 依然受之后 Keywords 逻辑控制。
            # 简化版：只要在 in_abstract 状态下且未检测到 KEYWORDS，默认延续
            current_section = SectionType.ABSTRACT
        elif not in_body and not in_references:
            # 头部启发式逻辑 (WJG)
            if is_wjg and _is_wjg_meta(text):
                current_section = SectionType.ARTICLE_INFO
            elif detected != SectionType.UNKNOWN:
                current_section = detected
            elif not found_title and not _is_affiliation_line(text) and not _is_author_line(text):
                current_section = SectionType.TITLE
                found_title = True
            elif found_title and not found_running_title and not _is_affiliation_line(text) and not _is_author_line(text):
                if len(text.split()) <= 10:
                    current_section = SectionType.RUNNING_TITLE
                    found_running_title = True
                else:
                    current_section = SectionType.TITLE
            elif not found_authors and _is_author_line(text):
                current_section = SectionType.AUTHORS
                found_authors = True
            elif (found_authors or found_affiliations) and _is_affiliation_line(text):
                current_section = SectionType.AFFILIATIONS
                found_affiliations = True
            else:
                # 检查是否进入正文 (非摘要模式下)
                if not in_abstract:
                    for pat in BODY_START_KEYWORDS:
                        if re.search(pat, text, re.IGNORECASE):
                            in_body = True
                            current_section = SectionType.BODY
                            break
        
        # 参考文献检测
        if in_body:
            for pat in BODY_END_KEYWORDS:
                if re.search(pat, text, re.IGNORECASE):
                    in_body = False
                    in_references = True
                    current_section = SectionType.REFERENCES
                    break
        
        # 如果 detected 为具体类型且当前不是正文/参考文献，更新它
        if detected != SectionType.UNKNOWN and not in_body and not in_references and not in_abstract:
             current_section = detected

        # 特殊处理：如果被识别为强正文关键词（如 Introduction），强制进入正文
        # 即使在 in_abstract 模式下，如果遇到明确的正文标题且不是第一段，也应考虑切出
        if not in_references:
             for pat in BODY_START_KEYWORDS:
                if re.search(pat, text, re.IGNORECASE):
                    # 如果已经在摘要中，只有当检测到的是强正文标志（如 Introduction）时才退出
                    # 或者，我们可以检查这是否是摘要的后续段落
                    if in_abstract and not re.search(r"^introduction\s*$", text, re.IGNORECASE):
                        continue
                        
                    in_body = True
                    in_abstract = False
                    current_section = SectionType.BODY
                    break
        para_labels.append((i, current_section, para))

    # Phase 2: 合并
    for idx, st, para in para_labels:
        if st not in ds.sections:
            ds.sections[st] = Section(section_type=st, start_index=idx)
        section = ds.sections[st]
        section.paragraphs.append((idx, para))
        section.end_index = idx

    for section in ds.sections.values():
        section.raw_text = section.text

    # Phase 3: 统计
    ref_section = ds.get_section(SectionType.REFERENCES)
    if ref_section:
        ref_count = 0
        for _, para in ref_section.paragraphs:
            t = para.text.strip()
            if t and re.match(r"^\d+\s", t):
                ref_count += 1
        ds.reference_count = ref_count

    return ds
