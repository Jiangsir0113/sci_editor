import sys
import re
from docx import Document
from sci_editor.parser import parse_document
from sci_editor.engine import RuleEngine
from sci_editor.models import SectionType

# Ensure UTF-8 output
sys.stdout.reconfigure(encoding='utf-8')

def test_new_fixes():
    doc = Document()
    # Paragraph 0: Article Info / Meta
    doc.add_paragraph("Manuscript Type: ORIGINAL ARTICLE")
    
    # Paragraph 1: Title
    doc.add_paragraph("New Study on Coffee")
    
    # Paragraph 2: Running Title
    doc.add_paragraph("Coffee Study")
    
    # Paragraph 3: Authors
    doc.add_paragraph("Yi Zeng, Pei-Yi Li")
    
    # Paragraph 4: Affiliations
    doc.add_paragraph("Department of Physics, university of Oxford, guangzhou 510000, China")
    
    # Paragraph 5: Abstract header
    doc.add_paragraph("Abstract")
    doc.add_paragraph("BACKGROUND: xxx. AIM: xxx. METHODS: xxx. RESULTS: xxx. CONCLUSION: xxx.")
    
    # Paragraph 7: Introduction (Body Start)
    doc.add_paragraph("Introduction")
    doc.add_paragraph("The results were compared using χ2 test (P < 0.05).")
    doc.add_paragraph("The area was 50 m2 and BMI was 24 kg/m2.")
    
    doc_path = "tests/test_new_refinements.docx"
    doc.save(doc_path)
    
    doc_struct = parse_document(doc_path)
    
    # Debug sections
    print("--- Section Debug ---")
    for st, section in doc_struct.sections.items():
        print(f"Section {st}: {len(section.paragraphs)} paragraphs")
    
    engine = RuleEngine()
    
    print("\n--- Running Check ---")
    issues = engine.check(doc_struct)
    for i in issues:
        print(f"[{i.rule_id}] {i.message} (Fixable: {i.fixable})")
        
    print("\n--- Running Fix ---")
    fixed_count = engine.fix_all(doc_struct, issues)
    print(f"Fixed {fixed_count} issues.")
    
    print("\n--- Final Text ---")
    # For affiliations:
    aff = doc_struct.get_section(SectionType.AFFILIATIONS)
    if aff:
        print(f"Affiliation: '{aff.paragraphs[0][1].text}'")
    
    # For body:
    body = doc_struct.get_section(SectionType.BODY)
    if body:
        for _, p in body.paragraphs:
            if "χ" in p.text or "m" in p.text:
                print(f"Body: '{p.text}'")

if __name__ == "__main__":
    test_new_fixes()
