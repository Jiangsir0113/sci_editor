from docx import Document
from sci_editor.parser import parse_document
from sci_editor.engine import RuleEngine
from sci_editor.models import SectionType

def test_user_broken_string():
    doc = Document()
    doc.add_paragraph("Abstract", style="Heading 1")
    # Using the exact problematic strings from user report
    text = (
        "After adjusting the influence of some other factors, we found that the chances of "
        "drinking coffee and diabetes are inversely related (OR = 0.79, 95%CI (0.64, 0.98, P = 0.029)), "
        "and the high level of 2h-PG is also inversely related (OR = 0.83, 95%CI (0.73, 0.95, P = 0.005)). "
        "For those who drink coffee with milk, the chance of high 2h-PG level is reduced by 24% "
        "(OR = 0.76, 95%CI (0.66, 0.88, P < 0.001)), and the chance of high HbA1c is reduced by 28% "
        "(OR = 0.72, 95%CI (0.63, 0.83, P < 0.001))."
    )
    p = doc.add_paragraph(text)
    
    doc_path = "tests/user_report.docx"
    doc.save(doc_path)
    
    doc_struct = parse_document(doc_path)
    engine = RuleEngine()
    
    issues = engine.check(doc_struct)
    print(f"Found {len(issues)} issues.")
    
    engine.fix_all(doc_struct, issues)
    
    text_after = doc_struct.all_paragraphs[1].text
    print("\n--- RESULT ---")
    print(text_after)
    
    # Check for missing parentheses
    if text_after.count("(") != text_after.count(")"):
        print(f"\nALERT: Balanced check failed - (:{text_after.count('(')} vs ):{text_after.count(')')}")
    
    # Check for italics
    print("\n--- ITALICS CHECK ---")
    for i, run in enumerate(doc_struct.all_paragraphs[1].runs):
        if run.text.strip() in ["P", "OR"]:
             print(f"Run {i}: '{run.text}' Italic: {run.italic}")

if __name__ == "__main__":
    test_user_broken_string()
