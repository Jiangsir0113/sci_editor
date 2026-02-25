import os
from docx import Document
from sci_editor.parser import parse_document
from sci_editor.models import SectionType
from sci_editor.engine import RuleEngine

def create_abbr_test_doc(filename):
    doc = Document()
    doc.add_paragraph("Manuscript Type: ORIGINAL ARTICLE")
    
    # 摘要部分：HCC 出现 2 次 -> 预期全部转为全称 Hepatocellular carcinoma
    doc.add_paragraph("Abstract")
    doc.add_paragraph("Hepatocellular carcinoma (HCC) is a common cancer. HCC treatment is difficult.")
    
    # 正文部分：OGTT 出现 4 次 -> 预期第一处定义，后续缩写
    doc.add_paragraph("Introduction")
    doc.add_paragraph("Oral glucose tolerance test (OGTT) is used for diagnosis.")
    doc.add_paragraph("We performed an OGTT on all patients.")
    doc.add_paragraph("The OGTT results were consistent.")
    doc.add_paragraph("Another OGTT was planned.")
    
    # 正文部分：HBV 出现 2 次 -> 预期全部转为全称 Hepatitis B virus
    doc.add_paragraph("Hepatitis B virus (HBV) is a risk factor. Patients with HBV need care.")
    
    doc.add_paragraph("Key Words: testing, abbr")
    doc.save(filename)

def test_abbreviation_enhanced():
    test_file = "tests/test_abbr_v2.docx"
    create_abbr_test_doc(test_file)
    
    print("--- Running Abbreviation Enhanced Test ---")
    struct = parse_document(test_file)
    engine = RuleEngine()
    
    # 检查
    issues = engine.check(struct)
    abbr_issues = [i for i in issues if i.rule_id == "3.10.x"]
    print(f"Detected {len(abbr_issues)} abbreviation issues.")
    for i in abbr_issues:
        print(f"  [{i.severity.value}] {i.message} (Context: {i.context})")
    
    # 修复
    fixed_count = engine.fix_all(struct, abbr_issues)
    print(f"Fixed {fixed_count} issues.")
    
    fixed_file = "tests/test_abbr_fixed_v2.docx"
    struct.word_doc.save(fixed_file)
    
    # 验证修复后的文本
    new_struct = parse_document(fixed_file)
    
    abs_text = new_struct.get_section_text(SectionType.ABSTRACT)
    print("\n[Abstract After Fix]")
    print(abs_text)
    # 预期不再有 "HCC"
    if "HCC" not in abs_text and "Hepatocellular carcinoma" in abs_text:
        print("PASS: Abstract Fix Success: Low freq term converted to full name.")
    else:
        print("FAIL: Abstract Fix Failed.")

    body_text = new_struct.get_section_text(SectionType.BODY)
    print("\n[Body After Fix]")
    print(body_text)
    # 预期 OGTT 第一处是定义，后续是缩写
    ogtt_defs = len(re.findall(r"Oral glucose tolerance test \(OGTT\)", body_text))
    ogtt_abbrs = len(re.findall(r"\bOGTT\b", body_text))
    if ogtt_defs == 1 and ogtt_abbrs == 4: # 4 includes the one in definition
        print("PASS: Body OGTT Success: High freq term defined once, then abbreviated.")
    else:
        print(f"FAIL: Body OGTT Failed: defs={ogtt_defs}, total_abbrs={ogtt_abbrs}")

    # HBV 预期全称
    if "HBV" not in body_text and "Hepatitis B virus" in body_text:
        print("PASS: Body HBV Success: Low freq term converted to full name.")
    else:
        print("FAIL: Body HBV Failed.")

import re
if __name__ == "__main__":
    if not os.path.exists("tests"):
        os.makedirs("tests")
    test_abbreviation_enhanced()
