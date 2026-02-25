"""Compare formatting between original and manually-edited documents."""
import sys
sys.path.insert(0, r'd:\AI_Project')
from docx import Document

doc_orig = Document(r'd:\AI_Project\tests\未修改.docx')
doc_edit = Document(r'd:\AI_Project\tests\人工修改.docx')

# Extract formatting details for ALL paragraphs in the EDITED version
print("=== EDITED DOCUMENT: Formatting details for first 25 paragraphs ===")
for i, p in enumerate(doc_edit.paragraphs[:25]):
    t = p.text.strip()
    if not t:
        continue
    print(f"\nP{i}: {t[:120]}")
    for j, r in enumerate(p.runs):
        fmt = []
        if r.bold: fmt.append('BOLD')
        if r.italic: fmt.append('ITALIC')
        if r.underline: fmt.append('UNDERLINE')
        if fmt and r.text.strip():
            print(f"  Run{j} [{', '.join(fmt)}]: '{r.text[:80]}'")

# Check headings in edited doc body
print("\n\n=== EDITED DOC: Body headings (all-caps and short paragraphs) ===")
for i, p in enumerate(doc_edit.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    # Short paragraph, all caps => heading
    letters = [c for c in t if c.isalpha()]
    if letters and all(c.isupper() for c in letters) and len(t.split()) <= 6 and len(t) < 50:
        runs_info = []
        for r in p.runs:
            f = []
            if r.bold: f.append('B')
            if r.italic: f.append('I')
            if r.underline: f.append('U')
            runs_info.append('/'.join(f) if f else '-')
        print(f"  P{i}: '{t}' => Runs: {runs_info}")

# Check italic in edited doc
print("\n\n=== EDITED DOC: Italic runs in body ===")
for i, p in enumerate(doc_edit.paragraphs):
    for r in p.runs:
        if r.italic and r.text.strip():
            txt = r.text.strip()
            if len(txt) < 30:
                print(f"  P{i}: italic '{txt}'")

# Check OR/CI formatting pattern in edited doc
print("\n\n=== EDITED DOC: OR/CI pattern samples ===")
for i, p in enumerate(doc_edit.paragraphs):
    t = p.text
    if '95%CI' in t or '95% CI' in t or 'OR' in t:
        # Show first occurrence
        if 'OR' in t and 'CI' in t and i > 20:
            print(f"  P{i}: {t[:200]}")
            break
