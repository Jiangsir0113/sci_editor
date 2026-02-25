import sys
from docx import Document
from docx.shared import Pt, RGBColor
import os
import io

# 强制 stdout 使用 utf-8 避免 Windows 控制台 GBK 错误
if sys.stdout.encoding.lower() != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# 确保能导入 sci_editor
sys.path.append(os.getcwd())

from sci_editor.parser import parse_document
from sci_editor.engine import RuleEngine

def test_formatting_preservation():
    # 1. 创建带特殊格式的测试文档
    doc = Document()
    
    # 机构行：测试首字母大写 + 不同字号和颜色
    p1 = doc.add_paragraph("Affiliations: ")
    p1.add_run("Southern Medical University, ").font.size = Pt(12)
    r_gz = p1.add_run("guangzhou")
    r_gz.font.size = Pt(14)
    r_gz.font.color.rgb = RGBColor(255, 0, 0) # 红色
    p1.add_run(" 510000, China")

    # 正文行：测试统计符号斜体 + 之前有特殊格式
    p2 = doc.add_paragraph("Results: ")
    r_bold = p2.add_run("Significant")
    r_bold.bold = True
    p2.add_run(" values were found (P < 0.05).")

    # 正文行：测试 χ2 上标 + 跨 Run 格式
    p3 = doc.add_paragraph("Analysis: ")
    p3.add_run("The data was analyzed using ")
    r_chi = p3.add_run("χ")
    r_chi.font.color.rgb = RGBColor(0, 0, 255) # 蓝色
    p3.add_run("2 test.")

    test_file = "tests/formatting_test.docx"
    doc.save(test_file)

    # 2. 运行工具修复
    ds = parse_document(test_file)
    engine = RuleEngine()
    issues = engine.check(ds)
    # 找出 fixable 的
    fixable = [i for i in issues if i.fixable]
    print(f"Found {len(fixable)} fixable issues.")
    
    engine.fix_all(ds, fixable)
    
    fixed_file = "tests/formatting_test_fixed.docx"
    ds.word_doc.save(fixed_file)

    # 3. 验证格式
    doc_fixed = Document(fixed_file)
    
    print("\n--- Verification ---")
    
    # 检查广州
    p1_fixed = doc_fixed.paragraphs[0]
    found_gz = False
    for r in p1_fixed.runs:
        if "Guangzhou" in r.text:
            print(f"Checking 'Guangzhou': Text='{r.text}', Size={r.font.size.pt if r.font.size else 'Default'}, Color={r.font.color.rgb}")
            if r.font.size and r.font.size.pt == 14 and r.font.color.rgb == RGBColor(255, 0, 0):
                print("SUCCESS: Affiliation formatting preserved.")
            else:
                print("FAILURE: Affiliation formatting lost.")
            found_gz = True
    
    # 检查 P 斜体
    p2_fixed = doc_fixed.paragraphs[1]
    found_p = False
    for r in p2_fixed.runs:
        if r.text == "P":
            print(f"Checking 'P': Italic={r.italic}, Bold={r.bold}")
            # P 应该是斜体，但前面的 "Significant" 应该是粗体
            found_p = True
    for r in p2_fixed.runs:
        if "Significant" in r.text:
            if r.bold:
                print("SUCCESS: Paragraph bold formatting preserved.")
            else:
                print("FAILURE: Paragraph bold formatting lost.")

    # 检查 χ²
    p3_fixed = doc_fixed.paragraphs[2]
    for r in p3_fixed.runs:
        if "χ²" in r.text:
            print(f"Checking 'χ²': Text='{r.text}', Color={r.font.color.rgb}")
            if r.font.color.rgb == RGBColor(0, 0, 255):
                print("SUCCESS: Chi-square color formatting preserved.")
            else:
                print("FAILURE: Chi-square color formatting lost.")

if __name__ == "__main__":
    if not os.path.exists("tests"):
        os.makedirs("tests")
    test_formatting_preservation()
