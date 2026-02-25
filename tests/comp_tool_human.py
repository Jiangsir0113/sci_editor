import sys
import os
import re
from docx import Document
from sci_editor.parser import parse_document
from sci_editor.engine import RuleEngine
from sci_editor.models import SectionType
from difflib import SequenceMatcher

# Ensure UTF-8 output
sys.stdout.reconfigure(encoding='utf-8')

def get_clean_text(para):
    return para.text.strip()

def get_formatting(para):
    info = []
    for run in para.runs:
        if run.text.strip():
            info.append({
                'text': run.text.strip(),
                'italic': run.italic,
                'bold': run.bold
            })
    return info

def is_heading(para):
    text = para.text.strip()
    if not text: return False
    # Heuristic for section headers
    if text.isupper() and len(text.split()) < 10:
        return True
    if text in ["BACKGROUND", "AIM", "METHODS", "RESULTS", "CONCLUSION", "Statistical analysis"]:
        return True
    return False

def compare():
    orig_path = "tests/未修改.docx"
    doc_struct = parse_document(orig_path)
    engine = RuleEngine()
    issues = engine.check(doc_struct)
    engine.fix_all(doc_struct, issues)
    
    tool_doc = doc_struct.word_doc
    human_doc = Document("tests/人工修改.docx")
    
    tool_paras = [p for p in tool_doc.paragraphs if p.text.strip()]
    human_paras = [p for p in human_doc.paragraphs if p.text.strip()]
    
    print(f"Non-empty paras - Tool: {len(tool_paras)}, Human: {len(human_paras)}")
    
    # alignment using SequenceMatcher
    tool_texts = [p.text.strip() for p in tool_paras]
    human_texts = [p.text.strip() for p in human_paras]
    
    sm = SequenceMatcher(None, tool_texts, human_texts)
    
    differences = []
    
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            # Text is same, check formatting
            for k in range(i2 - i1):
                p_t = tool_paras[i1 + k]
                p_h = human_paras[j1 + k]
                
                if is_heading(p_t): continue
                
                f_t = get_formatting(p_t)
                f_h = get_formatting(p_h)
                
                # Compare italics/bold keywords
                t_italics = {r['text'] for r in f_t if r['italic']}
                h_italics = {r['text'] for r in f_h if r['italic']}
                
                if t_italics != h_italics:
                    differences.append({
                        'type': 'italics',
                        'para_idx': i1 + k,
                        'text': p_t.text[:100],
                        'tool': list(t_italics),
                        'human': list(h_italics)
                    })
        
        elif tag == 'replace':
            for k in range(max(i2-i1, j2-j1)):
                t_txt = tool_paras[i1+k].text if i1+k < i2 else ""
                h_txt = human_paras[j1+k].text if j1+k < j2 else ""
                
                # Check for minor differences (like case or punctuation)
                if t_txt and h_txt:
                    if is_heading(tool_paras[i1+k]): continue
                    differences.append({
                        'type': 'text_diff',
                        'para_idx': i1 + k,
                        'tool': t_txt,
                        'human': h_txt
                    })
        elif tag == 'delete':
            for k in range(i1, i2):
                differences.append({
                    'type': 'extra_in_tool',
                    'para_idx': k,
                    'text': tool_paras[k].text
                })
        elif tag == 'insert':
            for k in range(j1, j2):
                differences.append({
                    'type': 'missing_in_tool',
                    'human_idx': k,
                    'text': human_paras[k].text
                })

    print("\n--- DETAILED DIFFERENCES ---")
    for d in differences:
        if d['type'] == 'text_diff':
            print(f"[Text Diff] Tool: \"{d['tool']}\" vs Human: \"{d['human']}\"")
        elif d['type'] == 'italics':
            print(f"[Italics Diff] Text: \"{d['text']}...\"\n  Tool: {d['tool']}\n  Human: {d['human']}")
        elif d['type'] == 'extra_in_tool':
            print(f"[Extra in Tool] \"{d['text']}\"")
        elif d['type'] == 'missing_in_tool':
            print(f"[Missing in Tool] \"{d['text']}\"")
        print("-" * 40)

if __name__ == "__main__":
    compare()
