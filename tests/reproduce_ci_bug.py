from docx import Document
from sci_editor.parser import parse_document
from sci_editor.engine import RuleEngine
from sci_editor.models import SectionType

def reproduce_ci_parenthesis_issue():
    # Create a document with the specific problematic sentence
    doc = Document()
    doc.add_paragraph("Abstract", style="Heading 1")
    p = doc.add_paragraph("After adjusting the influence of some other factors, we found that the chances of drinking coffee and diabetes are inversely related (OR = 0.79, 95%CI (0.64, 0.98, P = 0.029)), and the high level of 2h-PG is also inversely related (OR = 0.83, 95%CI (0.73, 0.95, P = 0.005)).")
    
    # Save for parsing
    doc_path = "tests/test_ci_bug.docx"
    doc.save(doc_path)
    
    # Run the engine
    doc_struct = parse_document(doc_path)
    engine = RuleEngine()
    
    print("\n--- BEFORE FIX ---")
    print(doc_struct.all_paragraphs[1].text)
    
    issues = engine.check(doc_struct)
    print(f"\nFound {len(issues)} issues.")
    for i in issues:
        print(f"  [{i.rule_id}] {i.message}")
        
    engine.fix_all(doc_struct, issues)
    
    print("\n--- AFTER FIX ---")
    text_after = doc_struct.all_paragraphs[1].text
    print(text_after)
    
    # Verification of parenthesis
    if text_after.count('(') != text_after.count(')'):
        print(f"FAIL: Parenthesis mismatch! '('={text_after.count('(')}, ')'={text_after.count(')')}")
    else:
        print("SUCCESS: Parenthesis are balanced.")

    # Verification of italics
    print("\n--- RUN INSPECTION ---")
    for i, run in enumerate(doc_struct.all_paragraphs[1].runs):
        italic_str = "[ITALIC]" if run.italic else ""
        print(f"  Run {i}: '{run.text}' {italic_str}")

if __name__ == "__main__":
    reproduce_ci_parenthesis_issue()
