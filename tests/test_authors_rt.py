import sys
from docx import Document
import os
import io

# 强制 stdout 使用 utf-8
if sys.stdout.encoding.lower() != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

sys.path.append(os.getcwd())

from sci_editor.parser import parse_document
from sci_editor.engine import RuleEngine
from sci_editor.models import SectionType

def test_improvements():
    # 1. 创建测试文档
    doc = Document()
    doc.add_heading("Article Title", 0)
    doc.add_paragraph("Running Title: PTMs in HCC senescence")
    doc.add_paragraph("Tai-xian Song, Hao-min Ji, Xing-sheng Shu")
    doc.add_paragraph("The People's Hospital, University of SC")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("Background: Metabolism-linked PTMs driving immunosuppressive senescence.")
    doc.add_paragraph("Introduction")
    doc.add_paragraph("Hepatocellular carcinoma (HCC) is a major global health threat...")
    
    test_file = "tests/reproduce_rt_author.docx"
    doc.save(test_file)
    
    # 2. 解析文档
    structure = parse_document(test_file)
    
    # 打印解析出的作者和短标题
    print(f"Parsed Authors: {structure.get_section_text(SectionType.AUTHORS)}")
    print(f"Parsed RT: {structure.get_section_text(SectionType.RUNNING_TITLE)}")
    
    # 3. 运行规则引擎
    engine = RuleEngine()
    print(f"Loaded rules: {engine.get_rule_ids()}")
    issues = engine.check(structure)
    
    print(f"\nFound {len(issues)} issues:")
    target_issues = []
    for issue in issues:
        if issue.rule_id in ["3.3.1", "3.2.1"]:
            print(f"[{issue.rule_id}] {issue.message} | Context: {issue.context} | Fixable: {issue.fixable}")
            target_issues.append(issue)
    
    # 4. 执行修复
    print("\nStarting fix_all...")
    fixed_count = engine.fix_all(structure, target_issues)
    print(f"\nFixed {fixed_count} issues.")
    
    # 5. 验证结果
    structure.word_doc.save("tests/reproduce_rt_author_fixed.docx")
    
    # 重新解析验证
    new_structure = parse_document("tests/reproduce_rt_author_fixed.docx")
    print(f"\nFixed Authors: {new_structure.get_section_text(SectionType.AUTHORS).strip()}")
    print(f"Fixed RT: {new_structure.get_section_text(SectionType.RUNNING_TITLE).strip()}")
    
    # 检查 RT 是否包含前缀且 et al 是否为斜体
    rt_section = new_structure.get_section(SectionType.RUNNING_TITLE)
    if rt_section and rt_section.paragraphs:
        rt_para = rt_section.paragraphs[0][1]
        print(f"RT Text: {rt_para.text}")
        for run in rt_para.runs:
             # 输出 run 的文本和斜体状态
             italic_str = "斜体" if run.italic else "常规"
             print(f"  Run: '{run.text}' [{italic_str}]")

if __name__ == "__main__":
    test_improvements()
