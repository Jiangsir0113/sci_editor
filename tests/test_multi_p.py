from docx import Document
from sci_editor.parser import parse_document
from sci_editor.engine import RuleEngine
from sci_editor.models import SectionType

def test_multiple_p_symbols():
    doc = Document()
    doc.add_paragraph("Abstract", style="Heading 1")
    
    # Sentence with multiple non-italic symbols
    text = "(OR = 0.79, 95%CI (0.64, 0.98, P = 0.029)), and (OR = 0.83, 95%CI (0.73, 0.95, P = 0.005))."
    p = doc.add_paragraph(text)
    
    doc_path = "tests/test_multi_p.docx"
    doc.save(doc_path)
    
    doc_struct = parse_document(doc_path)
    engine = RuleEngine()
    
    issues = engine.check(doc_struct)
    print(f"Found {len(issues)} issues.")
    
    # Fix all - this will trigger CI fix (which re-italizes) AND StatItalic fix
    engine.fix_all(doc_struct, issues)
    
    para = doc_struct.all_paragraphs[1]
    print("\n--- FINAL TEXT ---")
    print(para.text)
    
    if para.text.count("(") != para.text.count(")"):
        print("!! MISMATCH !!")
        
    print("\n--- RUNS ---")
    for i, run in enumerate(para.runs):
        print(f"Run {i}: '{run.text}' Italic: {run.italic}")

if __name__ == "__main__":
    test_multiple_p_symbols()
