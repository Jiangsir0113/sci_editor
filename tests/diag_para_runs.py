import sys
from docx import Document
import re

sys.stdout.reconfigure(encoding='utf-8')

def diag():
    doc = Document('tests/未修改.docx')
    # Paragraph 103 (0-indexed) seems to be the one with chi-square
    p = doc.paragraphs[103]
    print(f"Paragraph 103 Text: {p.text}")
    print(f"Paragraph 103 Chars: {[ord(c) for c in p.text]}")
    print("Runs:")
    for r in p.runs:
        print(f" - '{r.text}' (Italic={r.italic}, Superscript={r.font.superscript})")
        print(f"   Chars: {[ord(c) for c in r.text]}")

    # Let's also check affiliation at around index 11
    p11 = doc.paragraphs[11]
    print(f"\nParagraph 11 Text: {p11.text}")
    print(f"Paragraph 11 Chars: {[ord(c) for c in p11.text]}")

if __name__ == "__main__":
    diag()
