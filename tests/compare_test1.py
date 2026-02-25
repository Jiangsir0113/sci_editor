
import sys
import os
from docx import Document
import difflib

def get_doc_structure(path):
    if not os.path.exists(path):
        return f"File not found: {path}"
    
    doc = Document(path)
    lines = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        # 记录文本和一些基本格式信息
        fmt_info = []
        for run in para.runs:
            if run.bold: fmt_info.append("B")
            if run.italic: fmt_info.append("I")
            if run.font.color and run.font.color.rgb:
                fmt_info.append(f"C({run.font.color.rgb})")
        
        lines.append(f"P{i}: {text} [{'|'.join(set(fmt_info))}]")
    
    # 也检查表格
    for i, table in enumerate(doc.tables):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                lines.append(f"T{i}R{r}C{c}: {cell.text.strip()}")
                
    return lines

def compare_docs(orig_path, human_path, tool_path):
    print(f"Comparing documents in {os.path.dirname(orig_path)}...")
    
    orig = get_doc_structure(orig_path)
    human = get_doc_structure(human_path)
    tool = get_doc_structure(tool_path)
    
    # 输出到文件以便查看
    with open("tests/test1_orig.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(orig))
    with open("tests/test1_human.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(human))
    with open("tests/test1_tool.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(tool))
        
    print("Extracted structures to tests/test1_*.txt")
    
    # 对比 Human vs Tool
    diff = difflib.unified_diff(tool, human, fromfile='Tool', tofile='Human', lineterm='')
    
    with open("tests/test1_diff_tool_human.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(list(diff)))
    
    print("Diff (Tool -> Human) saved to tests/test1_diff_tool_human.txt")

if __name__ == "__main__":
    base_dir = "d:/AI_Project/test1"
    compare_docs(
        os.path.join(base_dir, "未修改.docx"),
        os.path.join(base_dir, "人工修改.docx"),
        os.path.join(base_dir, "工具修改.docx")
    )
