from docx import Document
from sci_editor.parser import parse_document
from sci_editor.engine import RuleEngine
from sci_editor.models import SectionType

def simulate_bug():
    doc = Document()
    doc.add_paragraph("Abstract", style="Heading 1")
    
    # Case 1: Complex nested parens
    p1 = doc.add_paragraph()
    p1.add_run("After adjusting factors, we found (OR = 0.83, 95%CI (0.73, 0.95, P = 0.005)).")
    
    # Case 2: Split runs - very likely cause
    p2 = doc.add_paragraph()
    p2.add_run("Another result reduced by 28% (OR = 0.72, 95%CI (")
    p2.add_run("0.63, 0.83, P < 0.001)).")
    
    # Case 3: Space variations
    p3 = doc.add_paragraph("Check this: (OR=0.76, 95%CI:(0.66, 0.88, P<0.001)).")

    doc_path = "tests/sim_bug.docx"
    doc.save(doc_path)
    
    doc_struct = parse_document(doc_path)
    engine = RuleEngine()
    
    issues = engine.check(doc_struct)
    engine.fix_all(doc_struct, issues)
    
    print("\n--- RESULTS ---")
    for i, p in enumerate(doc_struct.all_paragraphs):
        if i >= 1: # skip abstract
            print(f"Para {i}: '{p.text}'")
            if p.text.count("(") != p.text.count(")"):
                print("  !! PARENTHESIS MISMATCH !!")
            
            # Check italics
            p_italic = any(run.text == "P" and run.italic for run in p.runs)
            print(f"  P italicized: {p_italic}")

if __name__ == "__main__":
    simulate_bug()
