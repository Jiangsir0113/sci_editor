import os
from docx import Document
from sci_editor.parser import parse_document
from sci_editor.models import SectionType
from sci_editor.engine import RuleEngine

def create_test_doc(filename, manuscript_type, abstract_content):
    doc = Document()
    doc.add_paragraph(f"Manuscript Type: {manuscript_type}")
    doc.add_paragraph("Abstract")
    for p in abstract_content:
        doc.add_paragraph(p)
    doc.add_paragraph("Key Words: testing, rules, editor")
    doc.add_paragraph("Introduction")
    doc.add_paragraph("Hepatocellular carcinoma (HCC) is a major threat. OR 0.85 (95%CI (0.7, 1.2), P=0.01).")
    doc.save(filename)

def test_abstract_ci_or():
    print("=== Testing Review Type Abstract ===")
    review_file = "tests/test_review.docx"
    create_test_doc(review_file, "Review", ["This is a monolithic abstract without headings.", "It has two paragraphs though."])
    
    struct = parse_document(review_file)
    print(f"Manuscript Type: {struct.manuscript_type}")
    abs_section = struct.get_section(SectionType.ABSTRACT)
    print(f"Abstract para count: {len(abs_section.paragraphs)}")
    
    engine = RuleEngine()
    issues = engine.check(struct)
    for issue in issues:
        if issue.rule_id == "3.9":
            print(f"Issue [3.9]: {issue.message}")

    print("\n=== Testing Original Type Abstract (Success Case) ===")
    success_file = "tests/test_original_success.docx"
    create_test_doc(success_file, "Original Article", [
        "BACKGROUND: Background text.",
        "AIM: Aim text.",
        "METHODS: Methods text.",
        "RESULTS: Results text.",
        "CONCLUSION: Conclusion text."
    ])
    struct = parse_document(success_file)
    issues = engine.check(struct)
    abstract_issues = [i for i in issues if i.rule_id == "3.9"]
    if not abstract_issues:
        print("Success: No abstract structure issues found for complete original article.")
    else:
        for issue in abstract_issues:
            print(f"FAILED: Found issue [3.9]: {issue.message}")

    print("\n=== Testing CI/OR Fix ===")
    # Using the same original_file which has "OR 0.85 (95%CI (0.7, 1.2), P=0.01)."
    target_issues = [i for i in issues if i.rule_id in ["3.15.3", "3.15.4"]]
    print(f"Found {len(target_issues)} CI/OR issues.")
    
    fixed_count = engine.fix_all(struct, target_issues)
    print(f"Fixed {fixed_count} CI/OR issues.")
    
    fixed_file = "tests/test_fix_result.docx"
    struct.word_doc.save(fixed_file)
    
    new_struct = parse_document(fixed_file)
    body_text = new_struct.get_section_text(SectionType.BODY)
    print(f"Body Text after fix: {body_text.strip()}")

if __name__ == "__main__":
    if not os.path.exists("tests"):
        os.makedirs("tests")
    test_abstract_ci_or()
