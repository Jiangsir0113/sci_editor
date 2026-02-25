from docx import Document
from sci_editor.parser import parse_document
from sci_editor.engine import RuleEngine
from sci_editor.models import SectionType

def reproduce():
    # Create a dummy doc with headers to trigger section parsing
    doc = Document()
    doc.add_paragraph("Abstract", style="Heading 1")
    p1 = doc.add_paragraph("Summary of results.")
    
    doc.add_paragraph("Methods", style="Heading 1")
    # This should trigger StatisticsItalicRule if it has non-italic P
    p2 = doc.add_paragraph("Then we found (OR = 0.83, 95%CI (0.73, 0.95), P = 0.005) and more.")
    # This should trigger CIFormatRule
    p3 = doc.add_paragraph("Another result: (OR = 0.72, 95%CI (0.63, 0.83, P < 0.001)).")
    
    doc.save("tests/reproduce_bug.docx")
    
    doc_struct = parse_document("tests/reproduce_bug.docx")
    engine = RuleEngine()
    
    print("--- Before Fix ---")
    for s in doc_struct.sections.values():
        print(f"Section {s.section_type}:")
        for _, p in s.paragraphs:
            print(f"  - {p.text}")
    
    issues = engine.check(doc_struct)
    print(f"\nFound {len(issues)} issues.")
    for i in issues:
        print(f"  - {i.rule_id} [{i.rule_name}]: {i.message} (Fixable: {i.fixable})")

    # Fix them
    print("\n--- Applying Fixes ---")
    fixed_count = engine.fix_all(doc_struct, issues)
    print(f"Fixed {fixed_count} issues.")
    
    print("\n--- After Fix ---")
    # Refresh doc_struct or just check paragraphs
    for p in doc_struct.all_paragraphs:
        if p.text.strip():
            print(f"  - '{p.text}'")
            
    # Check if text was moved to the end or broken
    for p in doc_struct.all_paragraphs:
        if "and more" in p.text and not p.text.startswith("Then we found"):
            print("FAIL: 'and more' was detached from its sentence!")

if __name__ == "__main__":
    reproduce()
